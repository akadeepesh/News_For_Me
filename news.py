import requests
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup

NEWS_API_KEY = "..."


def fetch_current_affairs():
    url = f"https://newsapi.org/v2/top-headlines"
    params = {
        "apiKey": NEWS_API_KEY,
        "sources": "the-hindu",
        "pageSize": 10,
        "show": "description",
    }
    response = requests.get(url, params=params)
    if response.status_code == 200:
        data = response.json()
        articles = data.get("articles", [])
    else:
        print("Failed to fetch current affairs.")
        return []
    return articles


def extract_article_content(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, "html.parser")
        article = soup.find("article")
        if article:
            text = article.get_text()
            return text
        else:
            print("Failed to find the article.")
            return ""
    else:
        print("Failed to fetch the webpage.")
        return ""


def generate_current_affairs_doc(articles):
    doc = Document()
    doc.add_heading(f"Current Affairs - {date.today().strftime('%Y-%m-%d')}", level=1)
    source_style = doc.styles.add_style("Source", 1)
    source_style.font.size = Pt(8)
    source_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for i, article in enumerate(articles[:10], 1):
        doc.add_heading(f"{i}. {article['title']}", level=2)
        doc.add_paragraph(article["description"])
        content = extract_article_content(article["url"])
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("Failed to fetch the article content.")
        doc.add_paragraph(f"Source: {article['source']['name']}", style="Source")
        doc.add_paragraph(f"URL: {article['url']}", style="Source")
        doc.add_paragraph()

    return doc


def main():
    articles = fetch_current_affairs()
    if articles:
        doc = generate_current_affairs_doc(articles)
        doc.save("current_affairs.docx")
        print("Daily current affairs document generated successfully.")
    else:
        print("No current affairs data found.")


if __name__ == "__main__":
    main()
